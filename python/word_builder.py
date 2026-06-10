"""
DB Report — Générateur de rapport Word (.docx) à partir du JSON de collecte.

Reproduit la structure d'un rapport mensuel type avec un design
soigné : page de garde, historique, diffusion, sommaire (TOC native Word),
synthèse, puis par serveur → système + bases.

Usage:
    python -m db_report.word.word_builder collecte.json rapport.docx
    python -m db_report.word.word_builder collecte.json rapport.docx --logo logo.png

Dépendance: python-docx (pip install python-docx)
"""
from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# --- Palette (bleu corporate sobre, lisible côté client) -------------------
NAVY        = RGBColor(0x1F, 0x33, 0x52)   # titres principaux
BLUE        = RGBColor(0x36, 0x5F, 0x91)   # titres secondaires
BLUE_HEADER = "4F81BD"                       # fond en-tête de tableau (hex sans #)
INK         = RGBColor(0x22, 0x22, 0x22)
INK_MUTED   = RGBColor(0x60, 0x5E, 0x5C)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)

# Verdict → couleur de fond cellule (hex)
VERDICT_FILL = {
    "OK":       "E2EFDA",   # vert clair
    "INFO":     "E2EFDA",
    "NOTICE":   "DEEAF6",   # bleu clair
    "WARNING":  "FCE4D6",   # orange clair
    "LICENSE":  "FCE4D6",
    "CRITICAL": "F8CBAD",   # rouge-orange
    "FAILED":   "F8CBAD",
    "ERROR":    "F8CBAD",
}
ROW_ALT_FILL = "F2F6FB"     # fond alterné des lignes
MONO_FILL    = "F1ECE4"     # fond des blocs preformatted

FONT_SANS = "Calibri"
FONT_MONO = "Consolas"


def verdict_color(text: str) -> str | None:
    """Retourne la couleur de fond hex selon le préfixe du verdict, ou None."""
    if not text:
        return VERDICT_FILL["OK"]
    prefix = text.strip().split(":")[0].strip().upper()
    # Gérer les variantes "ACTIVE (standby)" etc.
    for key, fill in VERDICT_FILL.items():
        if prefix.startswith(key):
            return fill
    if prefix in ("ACTIVE", "NOT_CONFIGURED", "INSTALLED", "OK", "NO_GAP", "NOT_APPLICABLE"):
        return VERDICT_FILL["OK"]
    return None


# --- Helpers bas niveau ----------------------------------------------------

def _set_cell_background(cell, hex_color: str):
    """Applique une couleur de fond (shading) à une cellule de tableau.
    Dans tcPr, w:shd se place après w:tcW et w:gridSpan. On insère à la
    position conforme au schéma."""
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:shd"))
    if existing is not None:
        tc_pr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    # Ancre : après gridSpan ou tcW si présents, sinon en tête
    grid_span = tc_pr.find(qn("w:gridSpan"))
    tc_w = tc_pr.find(qn("w:tcW"))
    ref = grid_span if grid_span is not None else tc_w
    if ref is not None:
        ref.addnext(shd)
    else:
        tc_pr.insert(0, shd)


def _set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    """Marges internes de cellule (en twips). Les sous-éléments de tcMar
    utilisent w:left/w:right (pas start/end) et doivent suivre l'ordre
    top, left, bottom, right."""
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:tcMar"))
    if existing is not None:
        tc_pr.remove(existing)
    margins = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    # tcMar vient après shd dans tcPr ; on append simplement (shd déjà placé avant)
    tc_pr.append(margins)


def _no_table_borders_except_grid(table):
    """Applique le style built-in 'Table Grid' (bordures fines sur la grille).
    Plus robuste que de bricoler w:tblBorders à la main (ordre du schéma)."""
    try:
        table.style = "Table Grid"
    except KeyError:
        # Si le style n'existe pas dans ce document, fallback manuel
        tbl_pr = table._tbl.tblPr
        if tbl_pr.find(qn("w:tblBorders")) is not None:
            return
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "CCCCCC")
            borders.append(el)
        tbl_pr.insert(0, borders)


def _add_toc(doc):
    """Insère un champ TOC natif Word (mis à jour à l'ouverture)."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = r'TOC \o "1-3" \h \z \u'
    fldChar_sep = OxmlElement("w:fldChar")
    fldChar_sep.set(qn("w:fldCharType"), "separate")
    fldChar_text = OxmlElement("w:t")
    fldChar_text.text = "Faites Ctrl+clic puis « Mettre à jour les champs » pour générer le sommaire."
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_sep)
    run._r.append(fldChar_text)
    run._r.append(fldChar_end)


def _enable_update_fields_on_open(doc):
    """Force Word à proposer la mise à jour des champs (TOC) à l'ouverture.
    Place w:updateFields juste après w:zoom (qui est le 1er élément), et
    s'assure que w:zoom a bien son attribut percent requis."""
    settings = doc.settings.element
    # 1. Réparer w:zoom si percent manquant
    zoom = settings.find(qn("w:zoom"))
    if zoom is not None and zoom.get(qn("w:percent")) is None:
        zoom.set(qn("w:percent"), "100")
    # 2. Ajouter updateFields s'il n'existe pas
    if settings.find(qn("w:updateFields")) is not None:
        return
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    # Dans CT_Settings, updateFields vient juste avant hdrShapeDefaults/compat.
    # compat est présent par défaut → on insère avant lui.
    compat = settings.find(qn("w:compat"))
    if compat is not None:
        compat.addprevious(update)
    else:
        settings.append(update)


# --- Construction des styles -----------------------------------------------

def _configure_styles(doc):
    """Configure les styles de base + headings pour un rendu propre."""
    normal = doc.styles["Normal"]
    normal.font.name = FONT_SANS
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK

    # Headings (outlineLevel via built-in => TOC fonctionne)
    h1 = doc.styles["Heading 1"]
    h1.font.name = FONT_SANS
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = NAVY

    h2 = doc.styles["Heading 2"]
    h2.font.name = FONT_SANS
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = BLUE

    h3 = doc.styles["Heading 3"]
    h3.font.name = FONT_SANS
    h3.font.size = Pt(12.5)
    h3.font.bold = True
    h3.font.color.rgb = BLUE


# --- Rendu des éléments ----------------------------------------------------

def _add_preformatted(doc, content: str):
    """Bloc monospace encadré pour les sorties shell/sqlplus."""
    if not content:
        return
    table = doc.add_table(rows=1, cols=1)
    _no_table_borders_except_grid(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _set_cell_background(cell, MONO_FILL)
    _set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
    cell.paragraphs[0].text = ""
    first = True
    for line in content.rstrip("\n").split("\n"):
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = FONT_MONO
        run.font.size = Pt(8.5)
        run.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _add_data_table(doc, columns, rows, verdict_column=None):
    """Tableau de données avec en-tête bleu, lignes alternées, verdicts colorés."""
    if not columns:
        return
    n_cols = len(columns)
    table = doc.add_table(rows=1, cols=n_cols)
    _no_table_borders_except_grid(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # En-tête
    hdr = table.rows[0]
    for j, col in enumerate(columns):
        cell = hdr.cells[j]
        _set_cell_background(cell, BLUE_HEADER)
        _set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(str(col))
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        run.font.name = FONT_SANS

    # Données
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        row_fill = ROW_ALT_FILL if i % 2 == 1 else None
        for j in range(n_cols):
            val = str(row[j]) if j < len(row) else ""
            cell = cells[j]
            _set_cell_margins(cell)
            # Couleur : verdict prioritaire, sinon alternance
            fill = None
            if verdict_column is not None and j == verdict_column:
                fill = verdict_color(val)
            if fill is None:
                fill = row_fill
            if fill:
                _set_cell_background(cell, fill)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(val)
            run.font.size = Pt(8.5)
            run.font.name = FONT_SANS
            # Verdict en gras
            if verdict_column is not None and j == verdict_column and val.strip():
                run.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _add_note(doc, note: str):
    """Annotation DBA sous une section (italique, encadré gauche)."""
    if not note:
        return
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run("\u2192 " + note)
    run.font.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = INK_MUTED


def _render_section(doc, section: dict):
    """Rend une section de base selon son type."""
    title = section.get("title", section.get("id", "Section"))
    doc.add_heading(title, level=3)

    stype = section.get("type", "preformatted")
    if stype == "preformatted":
        _add_preformatted(doc, section.get("content", ""))
    elif stype == "table":
        _add_data_table(
            doc,
            section.get("columns", []),
            section.get("rows", []),
            section.get("verdict_column"),
        )
    elif stype == "keyvalue":
        rows = section.get("rows", [])
        _add_data_table(doc, ["Paramètre", "Valeur"], rows)
    elif stype == "text":
        for para in section.get("content", "").split("\n\n"):
            doc.add_paragraph(para)

    _add_note(doc, section.get("note", ""))


# --- Sections principales --------------------------------------------------

def _build_cover(doc, report, logo_path):
    """Page de garde : logo + titre centré."""
    if logo_path and Path(logo_path).exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run()
        run.add_picture(str(logo_path), width=Inches(2.4))

    # Espace vertical
    for _ in range(10):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(report.get("client_name", "CLIENT") + " — " + report.get("title", "Rapport"))
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = NAVY

    if report.get("subtitle"):
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(report["subtitle"])
        r2.font.size = Pt(14)
        r2.font.color.rgb = BLUE

    doc.add_page_break()


def _build_history_diffusion(doc, report):
    """Tableaux Historique + Diffusion du document."""
    doc.add_heading("Historique du document", level=2)
    _add_data_table(doc, ["Version", "Date", "Auteur", "Remarques"],
                    [[report.get("version", "1.0"),
                      report.get("generated_at", "")[:10],
                      report.get("author", ""),
                      "Première version"]])

    doc.add_heading("Diffusion du document", level=2)
    diffusion = report.get("diffusion", [])
    if diffusion:
        _add_data_table(doc, ["Nom", "Société", "Fonction", "Email"],
                        [[d.get("name", ""), d.get("company", ""),
                          d.get("role", ""), d.get("email", "")] for d in diffusion])
    doc.add_page_break()


def _build_synthesis(doc, synthesis):
    """Synthèse : contexte + périmètre + points d'attention."""
    doc.add_heading("Synthèse", level=1)
    doc.add_heading("Contexte", level=2)
    if synthesis.get("context"):
        doc.add_paragraph(synthesis["context"])

    scope = synthesis.get("scope", [])
    if scope:
        _add_data_table(doc, ["Serveur", "Base", "Statut", "Remarques"],
                        [[s.get("server", ""), s.get("base", ""),
                          s.get("status", ""), s.get("remark", "")] for s in scope])

    highlights = synthesis.get("highlights", [])
    if highlights:
        doc.add_heading("Points d'attention", level=2)
        # Grouper par scope
        current_scope = None
        for h in highlights:
            scope_name = h.get("scope", "global")
            if scope_name != current_scope and scope_name != "global":
                doc.add_heading(f"Base {scope_name}", level=3)
                current_scope = scope_name
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(h.get("text", ""))
    doc.add_page_break()


def _build_host(doc, host):
    """Un serveur : section système (liste générique) puis ses bases."""
    hostname = host.get("hostname", "serveur")
    doc.add_heading(f"Serveur {hostname}", level=1)

    # Nouveau format : system_sections (liste ordonnée de sections génériques)
    system_sections = host.get("system_sections")
    if system_sections:
        doc.add_heading("Configuration système", level=2)
        for section in system_sections:
            _render_section(doc, section)
    else:
        # Rétrocompat : ancien format system (dict à clés figées)
        system = host.get("system", {})
        sys_blocks = [
            ("Les instances en cours d'exécution", system.get("instances")),
            ("Les listeners en cours d'exécution", system.get("listeners")),
            ("Occupation des espaces disque", system.get("disk_usage")),
        ]
        for title, content in sys_blocks:
            if content:
                doc.add_heading(title, level=3)
                _add_preformatted(doc, content)
        if system.get("disk_usage_note"):
            _add_note(doc, system["disk_usage_note"])
        if system.get("storage_disks"):
            doc.add_heading("Stockage", level=2)
            doc.add_heading("Liste des disques disponibles", level=3)
            _add_preformatted(doc, system["storage_disks"])

    doc.add_page_break()

    # Bases du serveur
    for db in host.get("databases", []):
        sid = db.get("sid", "BASE")
        doc.add_heading(f"Base {sid}", level=2)
        for section in db.get("sections", []):
            _render_section(doc, section)
        doc.add_page_break()


# --- Entrée principale -----------------------------------------------------

def build_report(data: dict, output_path: str, logo_path: str | None = None):
    doc = Document()

    # Page US Letter / marges 1 pouce (cohérent avec le HTML)
    section = doc.sections[0]
    section.page_width = Twips(11906)   # A4 (la version d'origine est en A4)
    section.page_height = Twips(16838)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    _configure_styles(doc)

    report = data.get("report", {})

    # En-tête de page
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run(f"{report.get('client_name','')} — {report.get('title','')}")
    hr.font.size = Pt(8)
    hr.font.color.rgb = INK_MUTED

    # Pied de page
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = fp.add_run(f"Version : {report.get('version','1.0')}, Date : {report.get('generated_at','')[:10]}")
    fr.font.size = Pt(8)
    fr.font.color.rgb = INK_MUTED

    # Contenu
    _build_cover(doc, report, logo_path)
    _build_history_diffusion(doc, report)

    doc.add_heading("Table des matières", level=1)
    _add_toc(doc)
    doc.add_page_break()

    if data.get("synthesis"):
        _build_synthesis(doc, data["synthesis"])

    for host in data.get("hosts", []):
        _build_host(doc, host)

    _enable_update_fields_on_open(doc)
    doc.save(output_path)
    return output_path


def main(argv):
    if len(argv) < 2:
        print("Usage: word_builder.py <collecte.json> <rapport.docx> [--logo logo.png]")
        return 2
    json_path = argv[0]
    out_path = argv[1]
    logo = None
    if "--logo" in argv:
        logo = argv[argv.index("--logo") + 1]
    with open(json_path, encoding="utf-8") as f:
        data = json.load(f)
    result = build_report(data, out_path, logo)
    print(f"Rapport Word généré : {result}")
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))
